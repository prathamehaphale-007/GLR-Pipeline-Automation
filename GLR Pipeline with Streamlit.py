import streamlit as st
import os
import tempfile
import json
import time
from io import BytesIO
from docx import Document
from groq import Groq
import fitz  

MODEL_NAME = "llama-3.3-70b-versatile"

try:
    import pythoncom
    from docx2pdf import convert
    PDF_CONVERSION_AVAILABLE = True
except ImportError:
    PDF_CONVERSION_AVAILABLE = False

def get_groq_client(api_key: str | None = None) -> Groq:
    if not api_key:
        raise RuntimeError("Groq API key not provided.")
    return Groq(api_key=api_key)

def extract_text_from_pdfs(pdf_files):
    """
    Extract text from ALL pages of ALL PDFs using PyMuPDF.
    """
    all_chunks = []
    for uploaded in pdf_files:
        try:
            pdf_bytes = uploaded.getvalue()
            if not pdf_bytes:
                continue

            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for i in range(len(doc)):
                page = doc.load_page(i)
                txt = page.get_text("text").strip()
                header = f"\n\n=== REPORT: {uploaded.name} | PAGE {i + 1} ===\n"
                all_chunks.append(header + txt)
            doc.close()
        except Exception as e:
            st.warning(f"Failed to read '{uploaded.name}': {e}")

    return "\n".join(all_chunks).strip()

def get_docx_text(doc: Document) -> str:
    """Flatten a DOCX into plain text (paragraphs + tables)."""
    chunks = []
    for p in doc.paragraphs:
        if p.text.strip():
            chunks.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    chunks.append(t)
    return "\n".join(chunks)

def make_docx_from_text(text: str) -> Document:
    """
    Create a new DOCX from plain text.
    """
    doc = Document()
    for block in text.split("\n\n"):
        lines = block.split("\n")
        if not any(line.strip() for line in lines):
            continue
        for line in lines:
            doc.add_paragraph(line)
        doc.add_paragraph("")
    return doc

EXTRACTION_FIELDS = [
    "INSURED_NAME", "INSURED_H_STREET", "INSURED_H_CITY", "INSURED_H_STATE", 
    "INSURED_H_ZIP", "DATE_LOSS", "DATE_INSPECTED", "DATE_RECEIVED", "TOL_CODE", 
    "MORTGAGEE", "MORTGAGE_CO", "CLAIM_NUMBER", "POLICY_NUMBER",
    "DWELLING_DESCRIPTION", "DWELLING_NARRATIVE", "OTHER_STRUCTURES_NARRATIVE", 
    "FENCING_NARRATIVE", "CONTENTS_NARRATIVE", "SUPPLEMENT_NARRATIVE", 
    "PRIORS_NARRATIVE", "CODE_ITEMS_NARRATIVE", "OP_NARRATIVE", "MICA_NARRATIVE", 
    "MORTGAGE_INFO_NARRATIVE", "CAUSE_ORIGIN_NARRATIVE", "SUBROGATION_NARRATIVE", 
    "SALVAGE_NARRATIVE",
]

def extract_structured_data_with_llm(client: Groq, report_text: str, model_name: str) -> dict:
    fields_list = ", ".join(EXTRACTION_FIELDS)
    system_prompt = (
        "You are an insurance-claims information-extraction engine.\n"
        "Your job is to build a structured JSON object capturing ALL relevant details.\n"
        f"You MUST output a single JSON object with exactly these string keys:\n{fields_list}\n\n"
        "STRICT RULES ABOUT VALUES:\n"
        "1. Every non-empty value MUST be an exact contiguous substring of the REPORT TEXT.\n"
        "2. If multiple candidates exist, choose the best one, but copy it exactly.\n"
        "3. NEVER return 'N/A' or invented values in the JSON. Use empty string \"\" if not found.\n"
    )
    user_prompt = (
        "### REPORT TEXT ###\n"
        f"{report_text}\n\n"
        "Now output ONLY the JSON object described in the instructions."
    )

    resp = client.chat.completions.create(
        model=model_name,
        temperature=0.0,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )
    content = resp.choices[0].message.content
    data = json.loads(content)
    
    cleaned: dict[str, str] = {}
    for k in EXTRACTION_FIELDS:
        v = data.get(k, "")
        cleaned[k] = str(v) if v is not None else ""
    return cleaned

def generate_report_from_template_with_llm(client: Groq, template_text: str, extracted: dict, model_name: str) -> str:
    extracted_json = json.dumps(extracted, indent=2)
    system_prompt = (
        "You are an experienced field adjuster writing a final 'General Loss Report'.\n"
        "Transform the TEMPLATE into a FINAL completed report using the JSON object.\n"
        "- Replace placeholders like [DATE_LOSS], [INSURED_NAME] with exact JSON values.\n"
        "- Use *_NARRATIVE fields to write the full narrative sections.\n"
        "- Do not ignore non-empty JSON fields.\n"
        "- Output plain text only."
    )
    user_prompt = (
        "### EXTRACTED DATA (JSON) ###\n"
        f"{extracted_json}\n\n"
        "### GENERAL LOSS REPORT TEMPLATE TEXT ###\n"
        f"{template_text}\n\n"
        "Return the FINAL completed General Loss Report as plain text."
    )

    resp = client.chat.completions.create(
        model=model_name,
        temperature=0.1,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )
    return resp.choices[0].message.content.strip()


st.set_page_config(page_title="GLR Final Polish", page_icon="🛡️", layout="wide")

st.title("🛡️ GLR Automation: Vision Edition (Groq)")
st.markdown(f"""
**System Status:** `Active`
**Model:** `{MODEL_NAME}`
**Objective:** Zero N/A, Perfect O&P Logic, Strict Extraction.
""")

with st.sidebar:
    st.header("Credentials")
    api_key = st.text_input("Groq API Key", type="password")
    
    st.divider()
    if PDF_CONVERSION_AVAILABLE:
        st.success("✅ PDF Conversion Active")
    else:
        st.warning("⚠️ PDF Conversion Inactive (Windows Only)")

col1, col2 = st.columns(2)

with col1:
    st.subheader("1. Template Document")
    template_file = st.file_uploader("Upload Empty Word Template (.docx)", type=["docx"])

with col2:
    st.subheader("2. Source Evidence")
    input_pdfs = st.file_uploader("Upload Forms & Photo Reports (.pdf)", type=["pdf"], accept_multiple_files=True)

if st.button("Generate Final Report", type="primary"):
    if not api_key or not template_file or not input_pdfs:
        st.error("Please provide API Key, Template, and Source Files.")
    else:
        status = st.status("Starting Analysis...", expanded=True)
        
        try:
            client = get_groq_client(api_key)
            temp_dir = tempfile.mkdtemp()

            status.write("📂 Extracting text from PDFs using PyMuPDF...")
            report_text = extract_text_from_pdfs(input_pdfs)
            
            if not report_text.strip():
                status.update(label="Failed", state="error")
                st.error("No text extracted from PDFs. Ensure they are text-based, not pure images.")
                st.stop()

            status.write("📄 Reading Template...")
            try:
                template_doc = Document(BytesIO(template_file.getvalue()))
                template_text = get_docx_text(template_doc)
            except Exception as e:
                status.update(label="Failed", state="error")
                st.error(f"Could not read DOCX template: {e}")
                st.stop()

            status.write("🧠 Extracting structured data (PennyMac, Dates, O&P)...")
            extracted_data = extract_structured_data_with_llm(client, report_text, MODEL_NAME)
            
            with st.expander("✅ Validated Extraction (View Logic)"):
                st.json(extracted_data)

            status.write("✏️ Writing Final Report...")
            final_report_text = generate_report_from_template_with_llm(client, template_text, extracted_data, MODEL_NAME)

            status.write("💾 Saving to DOCX...")
            out_doc = make_docx_from_text(final_report_text)
            
            out_docx_path = os.path.join(temp_dir, "Completed_GLR.docx")
            out_doc.save(out_docx_path)
            
            out_pdf_path = os.path.join(temp_dir, "Completed_GLR.pdf")
            pdf_ready = False
            if PDF_CONVERSION_AVAILABLE:
                status.write("📕 Converting to PDF...")
                pythoncom.CoInitialize()
                try:
                    convert(out_docx_path, out_pdf_path)
                    pdf_ready = True
                except Exception as e:
                    st.warning(f"PDF conversion failed: {e}")

            status.update(label="Analysis Complete!", state="complete", expanded=False)
            st.success("Report Generated! Check the extracted data specifically for 'PennyMac' and Dates.")

            d1, d2 = st.columns(2)
            
            with open(out_docx_path, "rb") as f:
                d1.download_button("⬇️ Download Word (.docx)", f, "Completed_GLR.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            if pdf_ready:
                with open(out_pdf_path, "rb") as f:
                    d2.download_button("⬇️ Download PDF (.pdf)", f, "Completed_GLR.pdf", mime="application/pdf")
            
        except Exception as e:
            st.error(f"Error during processing: {e}")